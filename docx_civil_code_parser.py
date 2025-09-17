#!/usr/bin/env python3
"""
Word文档民法典解析器
专门用于解析.docx格式的民法典文件
"""
import openai
import json
import re
import time
import os
from typing import List, Dict, Any, Optional
from docx import Document
from pathlib import Path

class DocxCivilCodeParser:
    """Word文档民法典解析器"""
    
    def __init__(self, api_key: str = None, model: str = "qwen-plus-latest"):
        """
        初始化解析器
        
        Args:
            api_key: OpenAI API密钥，如果为None则从环境变量获取
            model: 使用的模型名称，默认为qwen-plus-latest
        """
        self.api_key = api_key or "sk-a14dc6cb330d4061a8d4396461f166f1"
        self.model = model
        self.client = openai.OpenAI(
            api_key=self.api_key, 
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
            timeout=60.0  # 增加超时时间到60秒
        )
        
    def read_docx_file(self, file_path: str) -> str:
        """
        读取Word文档内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            文档内容字符串
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # 只添加非空段落
                    full_text.append(text)
            
            return '\n'.join(full_text)
        except FileNotFoundError:
            raise FileNotFoundError(f"文件 {file_path} 不存在")
        except Exception as e:
            raise Exception(f"读取Word文档时出错: {e}")
    
    def extract_articles_from_docx(self, file_path: str) -> List[Dict[str, str]]:
        """
        从Word文档中提取条款 - 改进版
        
        Args:
            file_path: Word文档路径
            
        Returns:
            条款列表，每个条款包含标题和内容
        """
        try:
            doc = Document(file_path)
            articles = []
            current_article = {"title": "", "content": ""}
            
            print(f"开始处理Word文档，共 {len(doc.paragraphs)} 个段落")
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检查是否是条款标题（包含"第"和"条"）
                if re.match(r'第[零一二三四五六七八九十百千\d]+条', text):
                    # 如果已有当前条款，先保存
                    if current_article["title"]:
                        articles.append(current_article.copy())
                        print(f"保存条款: {current_article['title'][:20]}... (内容长度: {len(current_article['content'])})")
                    
                    # 分离标题和内容
                    # 查找"第xxx条"后面的内容
                    match = re.match(r'(第[零一二三四五六七八九十百千\d]+条)\s*(.*)', text)
                    if match:
                        title = match.group(1)
                        content = match.group(2).strip()
                        current_article = {"title": title, "content": content}
                        print(f"开始新条款: {title}")
                        if content:
                            print(f"  标题行内容: {content[:50]}...")
                    else:
                        # 如果没有匹配到，说明只有标题
                        current_article = {"title": text, "content": ""}
                        print(f"开始新条款: {text}")
                else:
                    # 添加到当前条款内容
                    if current_article["title"]:
                        if current_article["content"]:
                            current_article["content"] += "\n" + text
                        else:
                            current_article["content"] = text
                        
                        # 每50个段落显示一次进度
                        if i % 50 == 0:
                            print(f"  段落 {i}: 添加内容 (当前内容长度: {len(current_article['content'])})")
            
            # 添加最后一个条款
            if current_article["title"]:
                articles.append(current_article)
                print(f"保存最后条款: {current_article['title'][:20]}... (内容长度: {len(current_article['content'])})")
            
            print(f"提取完成，共 {len(articles)} 条法规")
            return articles
            
        except Exception as e:
            raise Exception(f"提取条款时出错: {e}")
    
    def split_articles_by_regex(self, text: str) -> List[str]:
        """
        使用正则表达式分割条款
        
        Args:
            text: 文档文本
            
        Returns:
            分割后的条款列表
        """
        # 使用正则表达式找到所有"第xxx条"的位置，然后进行分割
        articles = re.split(r'(?=第[零一二三四五六七八九十百千\d]+条)', text)
        # 移除可能存在的空字符串
        articles = [art for art in articles if art.strip()]
        return articles
    
    def parse_single_article(self, article_text: str, article_index: int, document_context: str = "") -> Dict[str, Any]:
        """
        解析单个条款
        
        Args:
            article_text: 条款文本
            article_index: 条款索引
            document_context: 文档上下文信息
            
        Returns:
            解析结果字典
        """
        system_prompt = """
你是一个法律文书信息抽取助手。
你的任务是：根据输入的法律文件类型，提取关键信息，并转换为结构化 JSON。

请遵循以下规则：

1. **识别文件类型**：
- 如果是 **法律 / 法规 / 规章**，通常包含“第X条”，请逐条抽取 → 使用【条文型 JSON 模板】。
- 如果是 **通知 / 指导意见 / 部门解释**，通常是整篇文件 → 使用【文件型 JSON 模板】。
- 如果是 **司法解释**，可能逐条，也可能整篇 → 如果有“第X条”则用【条文型 JSON 模板】，否则用【文件型 JSON 模板】。
- 如果是 **判例 / 裁判文书**，请抽取当事人、案由、法院意见、裁判结果 → 使用【案例型 JSON 模板】。

2. **输出格式**：
- 严格输出 JSON（无多余文字）。
- 如果有多条（如多个条文），请放在 JSON 数组中。
- 如果是单篇（如通知/案例），可以只输出一个 JSON 对象。

3. **缺失字段**请填写 `null`，字段必须齐全。
4. `keywords` 提取该条文中的关键法律术语。
5. `summary` 用一句话总结条文核心规定。
6. `related_articles` 如果有提及其他条款，就列出来，否则空数组。
7. **条文内容**：每一条的“content”必须包含从该“第X条”开始，直到下一个“第Y条”之前的所有文字（包括条款内的段落、列举、前款/后款），不要遗漏或截断。
8.  *content* 内容要跟原文一样，不能加入自己的理解，直接复制原文内容即可。

---

### 【条文型 JSON 模板】
```json
{
"law_name": "",
"article_number": "",
"chapter": "",
"content": "",
"summary": "",
"keywords": [],
"scope": "",
"penalty": null,
"exceptions": null,
"related_articles": [],
"effective_date": "",
"amendment_date": "",
"validity_status": "",
 "document_number": "",
"legal_level": "",
"source_url": "",
"tags": [],
"jurisdiction": ""
}
````

### 【文件型 JSON 模板】（通知/解释类）

```json
{
"law_name": "",
"document_type": "",
"document_number": "",
"issuing_body": "",
"issue_date": "",
"effective_date": "",
"amendment_date": null,
"legal_level": "",
"jurisdiction": "",
"content": "",
"summary": "",
"keywords": [],
"scope": "",
"penalty": null,
"exceptions": null,
"related_documents": [],
"source_url": "",
"tags": []
}
```

### 【案例型 JSON 模板】（判例/裁判文书）

```json
{
"case_name": "",
"case_number": "",
"court": "",
"trial_date": "",
"document_type": "",
"legal_level": "裁判文书",
"jurisdiction": "",
"parties": {
"plaintiff": "",
"defendant": ""
},
"facts": "",
"claims": "",
"defenses": "",
"court_opinion": "",
"judgment": "",
"related_laws": [],
"summary": ""
}
```

---

### 示例输入（条文型）
```
《中华人民共和国企业所得税法实施条例》
时 效 性：现行有效
中华人民共和国国务院令 第512号

第二章 税务管理

第六十条　除国务院财政、税务主管部门另有规定外，固定资产计算折旧的最低年限如下：
　　（一）房屋、建筑物，为20年；
　　（二）飞机、火车、轮船、机器、机械和其他生产设备，为10年；
　　（三）与生产经营活动有关的器具、工具、家具等，为5年；
　　（四）飞机、火车、轮船以外的运输工具，为4年；
　　（五）电子设备，为3年。

第六十一条　从事开采石油、天然气等矿产资源的企业，在开始商业性生产前发生的费用和有关固定资产的折耗、折旧方法，由国务院财政、税务主管部门另行规定。

第六十二条　生产性生物资产按照以下方法确定计税基础：
　　（一）外购的生产性生物资产，以购买价款和支付的相关税费为计税基础；
　　（二）通过捐赠、投资、非货币性资产交换、债务重组等方式取得的生产性生物资产，以该资产的公允价值和支付的相关税费为计税基础。
　　前款所称生产性生物资产，是指企业为生产农产品、提供劳务或者出租等而持有的生物资产，包括经济林、薪炭林、产畜和役畜等。

```

### 示例输出 （条文型）

```json
[
  {
    "law_name": "中华人民共和国企业所得税法实施条例",
    "article_number": "第六十条",
    "chapter": null,
    "content": "除国务院财政、税务主管部门另有规定外，固定资产计算折旧的最低年限如下：\n　　（一）房屋、建筑物，为20年；\n　　（二）飞机、火车、轮船、机器、机械和其他生产设备，为10年；\n　　（三）与生产经营活动有关的器具、工具、家具等，为5年；\n　　（四）飞机、火车、轮船以外的运输工具，为4年；\n　　（五）电子设备，为3年。",
    "summary": "规定固定资产折旧的最低年限。",
    "keywords": ["固定资产", "折旧", "最低年限"],
    "scope": "企业纳税人",
    "penalty": null,
    "exceptions": null,
    "related_articles": ["第六十条"],
    "effective_date": null,
    "amendment_date": null,
   "document_number": "中华人民共和国国务院令 第512号",
    "validity_status": "现行有效",
    "legal_level": "法律",
    "source_url": "",
    "tags": ["固定资产", "折旧"],
    "jurisdiction": "全国"
  },
  {
    "law_name": "中华人民共和国企业所得税法实施条例",
    "article_number": "第六十一条",
    "chapter": null,
    "content": "从事开采石油、天然气等矿产资源的企业，在开始商业性生产前发生的费用和有关固定资产的折耗、折旧方法，由国务院财政、税务主管部门另行规定。",
    "summary": "明确开采石油、天然气企业固定资产折旧方法由国务院规定。",
    "keywords": ["石油", "天然气", "矿产资源", "折旧"],
    "scope": "开采石油天然气企业",
    "penalty": null,
    "exceptions": null,
    "related_articles": ["第六十一条"],
    "effective_date": null,
    "amendment_date": null,
    "validity_status": "现行有效",
    "document_number": "中华人民共和国国务院令 第512号",
    "legal_level": "法律",
    "source_url": "",
    "tags": ["矿产资源", "折旧"],
    "jurisdiction": "全国"
  }
]
```

"""

        # 构建你的 Prompt，明确指示输出格式
        user_prompt = f"""
请将以下民法典条款处理成 JSON 格式。

{document_context}

条款文本：
{article_text}

请直接输出 JSON，不要有其他任何解释。
"""
        print(document_context)
        
        try:
            # 调用 API
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0,
                response_format={"type": "json_object"}
            )

            # 提取模型回复
            model_response = response.choices[0].message.content
            print(f"  🤖 模型回复: {model_response}")

            # 尝试解析返回的 JSON
            json_output = json.loads(model_response)
            return json_output

        except json.JSONDecodeError as e:
            print(f"处理第 {article_index+1} 条时出错，返回的不是有效 JSON: {model_response}")
            return {"error": f"Failed to parse JSON for article {article_index+1}", "raw_response": model_response}
        except Exception as e:
            print(f"处理第 {article_index+1} 条时发生其他错误: {e}")
            return {"error": f"Other error for article {article_index+1}"}
    
    def parse_docx_civil_code(self, input_file: str, output_file: str = None, delay: float = 1.0, use_structured_extraction: bool = True) -> List[Dict[str, Any]]:
        """
        解析Word文档格式的民法典
        
        Args:
            input_file: 输入Word文档路径
            output_file: 输出JSON文件路径，如果为None则自动生成
            delay: API调用间隔（秒），避免速率限制
            use_structured_extraction: 是否使用结构化提取（推荐）
            
        Returns:
            所有解析结果的列表
        """
        if not output_file:
            input_path = Path(input_file)
            output_file = input_path.stem + "_parsed.json"
        
        print(f"开始读取Word文档: {input_file}")
        
        # 提取文档上下文信息
        document_context = self.extract_document_context(input_file)
        
        if use_structured_extraction:
            # 使用结构化提取方法
            print("使用结构化提取方法...")
            articles = self.extract_articles_from_docx(input_file)
            print(f"共找到 {len(articles)} 条法规")
            
            all_results = []
            for i, article in enumerate(articles):
                print(f"正在处理第 {i+1} / {len(articles)} 条: {article['title']}")
                print(f"  内容长度: {len(article['content'])} 字符")
                
                # 组合标题和内容
                full_text = f"{article['title']}\n{article['content']}"
                result = self.parse_single_article(full_text, i, document_context)
                all_results.append(result)
                
                # 添加延迟
                if delay > 0:
                    time.sleep(delay)
        else:
            # 使用传统文本分割方法
            print("使用传统文本分割方法...")
            full_text = self.read_docx_file(input_file)
            articles = self.split_articles_by_regex(full_text)
            print(f"共找到 {len(articles)} 条法规")
            
            all_results = []
            for i, article_text in enumerate(articles):
                print(f"正在处理第 {i+1} / {len(articles)} 条...")
                result = self.parse_single_article(article_text, i, document_context)
                all_results.append(result)
                
                # 添加延迟
                if delay > 0:
                    time.sleep(delay)
        
        # 保存结果
        print(f"正在保存结果到: {output_file}")
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(all_results, f, ensure_ascii=False, indent=4)
        
        print(f"处理完成！共处理 {len(all_results)} 条法规。结果已保存到 {output_file}")
        return all_results
    
    def preview_docx_content(self, file_path: str, max_articles: int = 5) -> None:
        """
        预览Word文档内容
        
        Args:
            file_path: Word文档路径
            max_articles: 最大预览条款数
        """
        try:
            print(f"预览Word文档: {file_path}")
            print("=" * 50)
            
            if self.extract_articles_from_docx(file_path):
                articles = self.extract_articles_from_docx(file_path)
                print(f"文档包含 {len(articles)} 条法规")
                print("\n前几条法规预览:")
                
                for i, article in enumerate(articles[:max_articles]):
                    print(f"\n{i+1}. {article['title']}")
                    print(f"   内容: {article['content'][:100]}...")
            else:
                # 如果没有找到结构化条款，显示原始文本
                text = self.read_docx_file(file_path)
                print("未找到结构化条款，显示原始文本预览:")
                print(text[:500] + "..." if len(text) > 500 else text)
                
        except Exception as e:
            print(f"预览文档时出错: {e}")
    
    def test_specific_article(self, file_path: str, article_number: str = "第三百一十一条") -> None:
        """
        测试特定条文的提取
        
        Args:
            file_path: Word文档路径
            article_number: 要测试的条文号
        """
        try:
            doc = Document(file_path)
            print(f"测试条文: {article_number}")
            print("=" * 50)
            
            found_article = False
            current_content = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检查是否是目标条款
                if text.startswith(article_number):
                    found_article = True
                    current_content = [text]
                    print(f"找到条文开始，段落 {i}: {text}")
                elif found_article:
                    # 检查是否是下一个条文
                    if re.match(r'第[零一二三四五六七八九十百千\d]+条', text):
                        print(f"条文结束，段落 {i}: {text}")
                        break
                    else:
                        current_content.append(text)
                        print(f"  段落 {i}: {text}")
            
            if found_article:
                print(f"\n完整条文内容:")
                print("-" * 30)
                full_content = "\n".join(current_content)
                print(full_content)
                print(f"\n内容长度: {len(full_content)} 字符")
            else:
                print(f"未找到条文: {article_number}")
                
        except Exception as e:
            print(f"测试时出错: {e}")
    
    def extract_document_context(self, file_path: str) -> str:
        """
        提取文档的上下文信息（标题、发文机关等）
        
        Args:
            file_path: Word文档路径
            
        Returns:
            文档上下文信息字符串
        """
        try:
            doc = Document(file_path)
            context_lines = []
            
            # 读取前几个段落，寻找文档信息
            for i, paragraph in enumerate(doc.paragraphs[:20]):  # 只检查前20个段落
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检查是否是文档标题或元信息
                # 使用更灵活的关键词匹配，考虑空格
                text_without_spaces = text.replace(" ", "")
                keywords = [
                    "中华人民共和国民法典", "发文机关", "发布日期", "生效日期", 
                    "时效性", "文号", "主席令", "全国人民代表大会"
                ]
                
                if any(keyword in text or keyword in text_without_spaces for keyword in keywords):
                    context_lines.append(text)
                
                # 如果遇到第一个条文，停止收集
                if re.match(r'第[零一二三四五六七八九十百千\d]+条', text):
                    break
            
            if context_lines:
                context = "\n".join(context_lines)
                print(f"📋 提取到文档上下文信息:")
                print(context)
                return context
            else:
                return ""
                
        except Exception as e:
            print(f"提取文档上下文时出错: {e}")
            return ""


def main():
    """主函数示例"""
    # 使用配置的API密钥
    api_key = "sk-a14dc6cb330d4061a8d4396461f166f1"  # 请替换为实际的API密钥
    model = "qwen-plus-latest"
    base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"
    
    print(f"✅ 使用配置的API密钥和模型: {model}")
    print(f"✅ 使用base_url: {base_url}")
    
    # 创建解析器实例
    parser = DocxCivilCodeParser(api_key=api_key, model=model)
    
    # 要解析的Word文档
    input_file = "1.中华人民共和国民法典.docx"
    
    # 检查文件是否存在
    if not os.path.exists(input_file):
        print(f"文件 {input_file} 不存在")
        print("请确保文件在当前目录中")
        return
    
    # 测试特定条文
    print("\n🔍 测试第三百一十一条的提取...")
    parser.test_specific_article(input_file, "第三百一十一条")
    
    # 预览文档内容
    parser.preview_docx_content(input_file)
    
    # 询问是否继续解析
    response = input("\n是否继续解析？(y/n): ")
    if response.lower() != 'y':
        print("解析已取消")
        return
    
    # 解析民法典文件
    try:
        results = parser.parse_docx_civil_code(
            input_file=input_file,
            output_file="民法典解析结果.json",
            delay=5,# 每次API调用间隔5秒
            use_structured_extraction=True  # 使用结构化提取
        )
        print(f"成功解析 {len(results)} 条法规")
    except Exception as e:
        print(f"解析过程中出现错误: {e}")


if __name__ == "__main__":
    main()
