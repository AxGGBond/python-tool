#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MySQL 数据导出到 Excel 工具
支持导出表数据、查询结果到 Excel 文件
"""

import sys
import argparse
from pathlib import Path
import subprocess
import json
import re

try:
    import pandas as pd
    import pymysql
    from sqlalchemy import create_engine
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print("缺少必要的库，请安装：")
    print("  pip install pandas pymysql openpyxl sqlalchemy python-docx")
    print(f"错误详情：{e}")
    sys.exit(1)


def clean_dataframe_for_excel(df):
    """
    清理 DataFrame 中的非法字符，使其可以正常导出到 Excel
    
    Excel/openpyxl 不支持的控制字符：
    - \x00-\x08: 控制字符
    - \x0B: 垂直制表符
    - \x0C: 换页符
    - \x0E-\x1F: 其他控制字符
    
    Args:
        df: pandas DataFrame
        
    Returns:
        清理后的 DataFrame
    """
    df_cleaned = df.copy()
    
    # 定义非法字符的正则表达式（Excel 不支持的控制字符）
    # 保留 \x09 (tab), \x0A (换行), \x0D (回车)
    illegal_chars_pattern = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F]')
    
    # 遍历所有列
    for col in df_cleaned.columns:
        if df_cleaned[col].dtype == 'object':  # 字符串类型
            # 清理字符串中的非法字符
            df_cleaned[col] = df_cleaned[col].astype(str).apply(
                lambda x: illegal_chars_pattern.sub('', x) if pd.notna(x) and x != 'nan' else x
            )
    
    return df_cleaned


class MySQLToExcelExporter:
    """MySQL 到 Excel 导出器"""
    
    def __init__(self, host="localhost", port=3306, user="root", password="", 
                 database="", use_docker=False, container_name="law-manage-mysql"):
        """
        初始化导出器
        
        Args:
            host: MySQL 主机地址
            port: MySQL 端口
            user: 用户名
            password: 密码
            database: 数据库名
            use_docker: 是否使用 Docker 容器
            container_name: Docker 容器名称
        """
        self.host = host
        self.port = port
        self.user = user
        self.password = password
        self.database = database
        self.use_docker = use_docker
        self.container_name = container_name
        self.connection = None
        self.engine = None
    
    def _get_engine(self):
        """获取 SQLAlchemy engine（用于 pandas）"""
        # 构建连接字符串
        if self.use_docker:
            # 使用 Docker 容器连接，通过端口映射
            host = 'localhost' if self.host == 'localhost' else self.host
        else:
            host = self.host
        
        # 对密码进行 URL 编码（处理特殊字符）
        from urllib.parse import quote_plus
        encoded_password = quote_plus(self.password) if self.password else ''
        
        # 构建 MySQL 连接字符串
        connection_string = f"mysql+pymysql://{self.user}:{encoded_password}@{host}:{self.port}/{self.database}?charset=utf8mb4"
        
        # 创建 SQLAlchemy engine
        return create_engine(connection_string, echo=False, pool_pre_ping=True)
    
    def _get_connection(self):
        """获取数据库连接（用于需要游标的操作）"""
        if self.use_docker:
            # 使用 Docker 容器连接
            host = 'localhost' if self.host == 'localhost' else self.host
        else:
            host = self.host
            
        return pymysql.connect(
            host=host,
            port=self.port,
            user=self.user,
            password=self.password,
            database=self.database,
            charset='utf8mb4',
            cursorclass=pymysql.cursors.DictCursor
        )
    
    def export_table(self, table_name, output_file=None, sheet_name=None, 
                     where_clause="", limit=None, batch_size=1000, batch_export=False):
        """
        导出表数据到 Excel
        
        Args:
            table_name: 表名
            output_file: 输出 Excel 文件路径（或输出目录，如果分批导出）
            sheet_name: Excel 工作表名称
            where_clause: WHERE 条件（不包含 WHERE 关键字）
            limit: 限制导出行数
            batch_size: 每批导出的行数（默认1000）
            batch_export: 是否分批导出（每批保存为一个文件）
        """
        try:
            self.connection = self._get_connection()
            self.engine = self._get_engine()
            
            # 先获取总行数
            count_query = f"SELECT COUNT(*) as total FROM `{table_name}`"
            if where_clause:
                count_query += f" WHERE {where_clause}"
            
            with self.connection.cursor() as cursor:
                cursor.execute(count_query)
                total_rows = cursor.fetchone()['total']
            
            if total_rows == 0:
                print(f"⚠ 表 {table_name} 没有数据")
                return False
            
            print(f"表 {table_name} 共有 {total_rows} 行数据")
            
            # 设置工作表名称
            if sheet_name is None:
                sheet_name = table_name[:31]  # Excel 工作表名称最长31字符
            
            if batch_export:
                # 分批导出
                return self._export_table_batch(
                    table_name, output_file, sheet_name, 
                    where_clause, limit, batch_size, total_rows
                )
            else:
                # 单文件导出（原有逻辑）
                # 构建查询语句
                query = f"SELECT * FROM `{table_name}`"
                if where_clause:
                    query += f" WHERE {where_clause}"
                if limit:
                    query += f" LIMIT {limit}"
                
                print(f"执行查询: {query}")
                
                # 优先使用 pymysql 直接读取（更可靠）
                try:
                    with self.connection.cursor() as cursor:
                        cursor.execute(query)
                        # 获取列名
                        columns = [desc[0] for desc in cursor.description]
                        # 获取数据
                        rows = cursor.fetchall()
                    
                    if rows:
                        # 转换为 DataFrame
                        df = pd.DataFrame(rows, columns=columns)
                    else:
                        df = pd.DataFrame(columns=columns)
                except Exception as e:
                    print(f"  使用直接连接失败: {e}，尝试使用 SQLAlchemy...")
                    # 备用方案：使用 SQLAlchemy
                    df = pd.read_sql(query, self.engine)
                
                if df.empty:
                    print(f"⚠ 表 {table_name} 没有数据")
                    return False
                
                # 调试：显示前几行数据预览
                print(f"  读取到 {len(df)} 行数据，列数: {len(df.columns)}")
                if len(df) > 0:
                    print(f"  前3行数据预览:")
                    for i in range(min(3, len(df))):
                        print(f"    行{i+1}: {dict(df.iloc[i].head(5))}")
                
                # 生成输出文件名
                if output_file is None:
                    output_file = f"{table_name}_export.xlsx"
                else:
                    output_file = Path(output_file)
                    # 如果指定的是目录，自动生成文件名
                    if output_file.is_dir() or (not output_file.suffix and output_file.exists() and output_file.is_dir()):
                        output_file = output_file / f"{table_name}_export.xlsx"
                    # 如果路径不存在但看起来像目录（没有扩展名），也当作目录处理
                    elif not output_file.suffix:
                        output_file = output_file / f"{table_name}_export.xlsx"
                    # 确保父目录存在
                    output_file.parent.mkdir(parents=True, exist_ok=True)
                    output_file = str(output_file)
                
                # 清理数据中的非法字符（处理 HTML 等内容）
                df = clean_dataframe_for_excel(df)
                
                # 导出到 Excel（确保第一行是列名）
                print(f"正在导出 {len(df)} 行数据到 {output_file}...")
                # index=False 确保不导出行索引，header=True 确保第一行是列名
                df.to_excel(output_file, index=False, sheet_name=sheet_name, engine='openpyxl', header=True)
                
                print(f"✓ 导出成功！")
                print(f"  文件: {output_file}")
                print(f"  行数: {len(df)}")
                print(f"  列数: {len(df.columns)}")
                
                return True
            
        except Exception as e:
            print(f"✗ 导出失败: {e}")
            import traceback
            traceback.print_exc()
            return False
        finally:
            if self.connection:
                self.connection.close()
    
    def _export_table_batch(self, table_name, output_file, sheet_name, 
                           where_clause, limit, batch_size, total_rows):
        """分批导出表数据"""
        try:
            # 确定输出目录
            if output_file is None:
                output_dir = Path(f"{table_name}_export")
            else:
                output_path = Path(output_file)
                if output_path.suffix == '.xlsx':
                    # 如果指定了文件，使用其所在目录
                    output_dir = output_path.parent / f"{table_name}_export"
                else:
                    # 如果指定了目录，直接使用
                    output_dir = output_path
            
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # 计算需要导出的批次数
            max_rows = limit if limit else total_rows
            num_batches = (max_rows + batch_size - 1) // batch_size  # 向上取整
            
            print(f"将分 {num_batches} 批导出，每批 {batch_size} 行")
            print(f"输出目录: {output_dir}")
            print("-" * 50)
            
            success_count = 0
            total_exported = 0
            
            for batch_num in range(num_batches):
                offset = batch_num * batch_size
                
                # 构建查询语句
                query = f"SELECT * FROM `{table_name}`"
                if where_clause:
                    query += f" WHERE {where_clause}"
                query += f" LIMIT {batch_size} OFFSET {offset}"
                
                # 读取当前批次数据（优先使用 pymysql 直接读取）
                try:
                    with self.connection.cursor() as cursor:
                        cursor.execute(query)
                        # 获取列名
                        columns = [desc[0] for desc in cursor.description]
                        # 获取数据
                        rows = cursor.fetchall()
                    
                    if rows:
                        # 转换为 DataFrame
                        df = pd.DataFrame(rows, columns=columns)
                    else:
                        df = pd.DataFrame(columns=columns)
                except Exception as e:
                    print(f"  使用直接连接失败: {e}，尝试使用 SQLAlchemy...")
                    # 备用方案：使用 SQLAlchemy
                    df = pd.read_sql(query, self.engine)
                
                if df.empty:
                    print(f"批次 {batch_num + 1}/{num_batches}: 没有数据，跳过")
                    break
                
                # 清理数据中的非法字符（处理 HTML 等内容）
                df = clean_dataframe_for_excel(df)
                
                # 生成输出文件名
                output_file = output_dir / f"{table_name}_{batch_num + 1}.xlsx"
                
                # 导出到 Excel（确保第一行是列名）
                print(f"批次 {batch_num + 1}/{num_batches}: 导出 {len(df)} 行 -> {output_file.name}")
                # index=False 确保不导出行索引，第一行自动是列名
                df.to_excel(output_file, index=False, sheet_name=sheet_name, engine='openpyxl', header=True)
                
                success_count += 1
                total_exported += len(df)
                
                # 如果当前批次数据少于 batch_size，说明已经是最后一批
                if len(df) < batch_size:
                    break
            
            print("-" * 50)
            print(f"✓ 分批导出完成！")
            print(f"  输出目录: {output_dir}")
            print(f"  成功批次: {success_count}/{num_batches}")
            print(f"  总行数: {total_exported}")
            print(f"  文件列表:")
            for i in range(1, success_count + 1):
                file_path = output_dir / f"{table_name}_{i}.xlsx"
                if file_path.exists():
                    file_size = file_path.stat().st_size / (1024 * 1024)  # MB
                    print(f"    - {file_path.name} ({file_size:.2f} MB)")
            
            return True
            
        except Exception as e:
            print(f"✗ 分批导出失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def export_query(self, query, output_file, sheet_name="Sheet1"):
        """
        导出查询结果到 Excel
        
        Args:
            query: SQL 查询语句
            output_file: 输出 Excel 文件路径
            sheet_name: Excel 工作表名称
        """
        try:
            self.connection = self._get_connection()
            self.engine = self._get_engine()
            
            print(f"执行查询: {query}")
            
            # 优先使用 pymysql 直接读取（更可靠）
            try:
                with self.connection.cursor() as cursor:
                    cursor.execute(query)
                    # 获取列名
                    columns = [desc[0] for desc in cursor.description]
                    # 获取数据
                    rows = cursor.fetchall()
                
                if rows:
                    # 转换为 DataFrame
                    df = pd.DataFrame(rows, columns=columns)
                else:
                    df = pd.DataFrame(columns=columns)
            except Exception as e:
                print(f"  使用直接连接失败: {e}，尝试使用 SQLAlchemy...")
                # 备用方案：使用 SQLAlchemy
                df = pd.read_sql(query, self.engine)
            
            if df.empty:
                print("⚠ 查询结果为空")
                return False
            
            # 处理输出文件路径
            output_file = Path(output_file)
            # 如果指定的是目录，自动生成文件名
            if output_file.is_dir() or (not output_file.suffix and output_file.exists() and output_file.is_dir()):
                output_file = output_file / "query_result.xlsx"
            # 如果路径不存在但看起来像目录（没有扩展名），也当作目录处理
            elif not output_file.suffix:
                output_file = output_file / "query_result.xlsx"
            # 确保父目录存在
            output_file.parent.mkdir(parents=True, exist_ok=True)
            output_file = str(output_file)
            
            # 清理数据中的非法字符（处理 HTML 等内容）
            df = clean_dataframe_for_excel(df)
            
            # 导出到 Excel（确保第一行是列名）
            print(f"正在导出 {len(df)} 行数据到 {output_file}...")
            # index=False 确保不导出行索引，header=True 确保第一行是列名
            df.to_excel(output_file, index=False, sheet_name=sheet_name, engine='openpyxl', header=True)
            
            print(f"✓ 导出成功！")
            print(f"  文件: {output_file}")
            print(f"  行数: {len(df)}")
            print(f"  列数: {len(df.columns)}")
            
            return True
            
        except Exception as e:
            print(f"✗ 导出失败: {e}")
            return False
        finally:
            if self.connection:
                self.connection.close()
    
    def export_all_tables(self, output_file=None, output_dir=None):
        """
        导出数据库中的所有表到 Excel（每个表一个工作表）
        
        Args:
            output_file: 输出 Excel 文件路径（所有表在一个文件中）
            output_dir: 输出目录（每个表一个文件）
        """
        try:
            self.connection = self._get_connection()
            self.engine = self._get_engine()
            
            # 获取所有表名
            with self.connection.cursor() as cursor:
                cursor.execute("SHOW TABLES")
                tables = [row[f'Tables_in_{self.database}'] for row in cursor.fetchall()]
            
            if not tables:
                print("数据库中没有表")
                return False
            
            print(f"找到 {len(tables)} 个表: {', '.join(tables)}")
            
            if output_file:
                # 导出到一个 Excel 文件，每个表一个工作表
                print(f"正在导出所有表到 {output_file}...")
                with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                    for table in tables:
                        try:
                            # 使用 pymysql 直接读取
                            query = f"SELECT * FROM `{table}`"
                            with self.connection.cursor() as cursor:
                                cursor.execute(query)
                                columns = [desc[0] for desc in cursor.description]
                                rows = cursor.fetchall()
                            
                            if rows:
                                df = pd.DataFrame(rows, columns=columns)
                            else:
                                df = pd.DataFrame(columns=columns)
                            
                            # 清理数据中的非法字符（处理 HTML 等内容）
                            df = clean_dataframe_for_excel(df)
                            
                            sheet_name = table[:31]  # Excel 工作表名称最长31字符
                            # index=False 确保不导出行索引，header=True 确保第一行是列名
                            df.to_excel(writer, index=False, sheet_name=sheet_name, header=True)
                            print(f"  ✓ {table}: {len(df)} 行")
                        except Exception as e:
                            print(f"  ✗ {table}: {e}")
                
                print(f"\n✓ 导出完成: {output_file}")
                return True
            
            elif output_dir:
                # 每个表导出为一个文件
                output_path = Path(output_dir)
                output_path.mkdir(parents=True, exist_ok=True)
                
                success_count = 0
                for table in tables:
                    try:
                        # 使用 pymysql 直接读取
                        query = f"SELECT * FROM `{table}`"
                        with self.connection.cursor() as cursor:
                            cursor.execute(query)
                            columns = [desc[0] for desc in cursor.description]
                            rows = cursor.fetchall()
                        
                        if rows:
                            df = pd.DataFrame(rows, columns=columns)
                        else:
                            df = pd.DataFrame(columns=columns)
                        
                        # 清理数据中的非法字符（处理 HTML 等内容）
                        df = clean_dataframe_for_excel(df)
                        
                        output_file = output_path / f"{table}.xlsx"
                        # index=False 确保不导出行索引，header=True 确保第一行是列名
                        df.to_excel(output_file, index=False, sheet_name=table[:31], engine='openpyxl', header=True)
                        print(f"  ✓ {table}: {len(df)} 行 -> {output_file}")
                        success_count += 1
                    except Exception as e:
                        print(f"  ✗ {table}: {e}")
                
                print(f"\n✓ 导出完成: {success_count}/{len(tables)} 个表")
                return True
            else:
                print("请指定 output_file 或 output_dir")
                return False
                
        except Exception as e:
            print(f"✗ 导出失败: {e}")
            return False
        finally:
            if self.connection:
                self.connection.close()
    
    def list_tables(self):
        """列出数据库中的所有表"""
        try:
            self.connection = self._get_connection()
            
            with self.connection.cursor() as cursor:
                cursor.execute("SHOW TABLES")
                tables = [row[f'Tables_in_{self.database}'] for row in cursor.fetchall()]
            
            if tables:
                print(f"数据库 '{self.database}' 中的表:")
                for i, table in enumerate(tables, 1):
                    # 获取表行数
                    cursor.execute(f"SELECT COUNT(*) as count FROM `{table}`")
                    count = cursor.fetchone()['count']
                    print(f"  {i}. {table} ({count} 行)")
            else:
                print("数据库中没有表")
            
            return tables
            
        except Exception as e:
            print(f"✗ 获取表列表失败: {e}")
            return []
        finally:
            if self.connection:
                self.connection.close()
    
    def export_law_regulation_to_word(self, output_dir=None, table_name="law_regulation"):
        """
        导出 law_regulation 表到 Word 文件（每行一个文件）
        
        Args:
            output_dir: 输出目录
            table_name: 表名（默认 law_regulation）
            
        Returns:
            bool: 是否成功
        """
        try:
            self.connection = self._get_connection()
            
            # 查询数据
            query = f"SELECT * FROM `{table_name}`"
            print(f"执行查询: {query}")
            
            with self.connection.cursor() as cursor:
                cursor.execute(query)
                # 由于使用了 DictCursor，fetchall() 返回的是字典列表
                rows = cursor.fetchall()
            
            if not rows:
                print(f"⚠ 表 {table_name} 没有数据")
                return False
            
            # 设置输出目录
            if output_dir is None:
                output_dir = Path(f"{table_name}_word_export")
            else:
                output_dir = Path(output_dir)
            
            output_dir.mkdir(parents=True, exist_ok=True)
            
            print(f"找到 {len(rows)} 条记录")
            print(f"输出目录: {output_dir}")
            
            # 调试：显示前3条记录的 title 值
            if rows:
                print("\n前3条记录的 title 字段预览:")
                for i, row_dict in enumerate(rows[:3], 1):
                    # rows 已经是字典列表，不需要转换
                    title_preview = row_dict.get('title', '（无title字段）')
                    print(f"  {i}. title = {repr(title_preview)[:80]}")
                    # 显示第一条记录的完整结构（调试用）
                    if i == 1:
                        print(f"      第一条记录的所有字段: {list(row_dict.keys())}")
                        print(f"      第一条记录的示例值: id={row_dict.get('id')}, title={row_dict.get('title')}")
            
            print("-" * 50)
            
            success_count = 0
            failed_count = 0
            
            for idx, row_dict in enumerate(rows, 1):
                try:
                    # rows 已经是字典列表（因为使用了 DictCursor），直接使用
                    
                    # 获取 title 作为文件名
                    title = row_dict.get('title', '')
                    if not title or str(title).strip() == '' or str(title).strip().lower() == 'title':
                        # 如果 title 为空或无效，尝试使用 id 或其他字段
                        record_id = row_dict.get('id', idx)
                        title = f'未命名_{record_id}'
                    
                    title_str = str(title).strip()
                    
                    # 清理文件名（移除非法字符）
                    safe_filename = re.sub(r'[<>:"/\\|?*]', '_', title_str)
                    safe_filename = safe_filename.strip()
                    
                    # 如果清理后为空，使用索引
                    if not safe_filename:
                        safe_filename = f'未命名_{idx}'
                    
                    # 限制长度，但保留足够的内容
                    if len(safe_filename) > 100:
                        safe_filename = safe_filename[:100]
                    
                    # 确保文件名唯一（如果已存在同名文件，添加序号）
                    base_filename = safe_filename
                    counter = 1
                    while (output_dir / f"{safe_filename}.docx").exists():
                        safe_filename = f"{base_filename}_{counter}"
                        counter += 1
                    
                    # 获取 content 作为正文
                    content = row_dict.get('content', '')
                    if content is None:
                        content = ''
                    
                    # 创建 Word 文档
                    doc = Document()
                    
                    # 设置标题样式
                    title_para = doc.add_heading(title, level=1)
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 添加其他字段作为元数据
                    doc.add_paragraph()  # 空行
                    metadata_para = doc.add_paragraph('【基本信息】')
                    metadata_para.runs[0].bold = True
                    metadata_para.runs[0].font.size = Pt(12)
                    
                    # 添加其他字段（除了 title 和 content）
                    for key, value in row_dict.items():
                        if key not in ['title', 'content']:
                            # 确保获取的是实际值，不是字段名
                            if value is not None:
                                value_str = str(value).strip()
                                # 如果值是字段名本身（说明数据有问题），跳过
                                if value_str.lower() == key.lower():
                                    continue
                                
                                if value_str != '':
                                    # 清理 HTML 标签（如果字段包含 HTML）
                                    value_str = re.sub(r'<[^>]+>', '', value_str)
                                    # 清理控制字符
                                    value_str = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', value_str)
                                    
                                    field_para = doc.add_paragraph(f"{key}: {value_str}")
                                    field_para.style.font.size = Pt(10)
                    
                    # 添加正文
                    doc.add_paragraph()  # 空行
                    content_para = doc.add_paragraph('【正文内容】')
                    content_para.runs[0].bold = True
                    content_para.runs[0].font.size = Pt(12)
                    
                    doc.add_paragraph()  # 空行
                    
                    # 处理 content（法规正文内容，可能包含 HTML）
                    if content:
                        content_str = str(content)
                        
                        # 移除 HTML 标签，保留文本内容
                        # 先处理一些常见的 HTML 标签，保留段落结构
                        content_str = re.sub(r'<br\s*/?>', '\n', content_str, flags=re.IGNORECASE)  # <br> 转为换行
                        content_str = re.sub(r'</p>', '\n\n', content_str, flags=re.IGNORECASE)  # </p> 转为双换行
                        content_str = re.sub(r'</div>', '\n', content_str, flags=re.IGNORECASE)  # </div> 转为换行
                        content_str = re.sub(r'<[^>]+>', '', content_str)  # 移除所有其他 HTML 标签
                        
                        # 清理控制字符（保留换行符和制表符）
                        content_str = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', content_str)
                        
                        # 清理多余的空白行（保留段落分隔）
                        content_str = re.sub(r'\n{3,}', '\n\n', content_str)
                        
                        # 按段落分割
                        paragraphs = content_str.split('\n\n')
                        
                        # 添加正文段落
                        for para_text in paragraphs:
                            para_text = para_text.strip()
                            if para_text:
                                # 如果段落很长，按单行换行分割
                                if '\n' in para_text:
                                    lines = para_text.split('\n')
                                    for line in lines:
                                        line = line.strip()
                                        if line:
                                            doc.add_paragraph(line)
                                else:
                                    doc.add_paragraph(para_text)
                    else:
                        doc.add_paragraph('（无正文内容）')
                    
                    # 保存文件
                    output_file = output_dir / f"{safe_filename}.docx"
                    doc.save(str(output_file))
                    
                    # 显示详细信息
                    title_display = title_str[:50] if len(title_str) > 50 else title_str
                    print(f"  ✓ [{idx}/{len(rows)}] 标题: {title_display}")
                    print(f"      文件名: {output_file.name}")
                    success_count += 1
                    
                except Exception as e:
                    print(f"  ✗ [{idx}/{len(rows)}] 导出失败: {e}")
                    failed_count += 1
                    import traceback
                    traceback.print_exc()
            
            print("-" * 50)
            print(f"✓ 导出完成！")
            print(f"  输出目录: {output_dir}")
            print(f"  成功: {success_count} 个文件")
            print(f"  失败: {failed_count} 个文件")
            print(f"  总计: {len(rows)} 条记录")
            
            return True
            
        except Exception as e:
            print(f"✗ 导出失败: {e}")
            import traceback
            traceback.print_exc()
            return False
        finally:
            if self.connection:
                self.connection.close()


def main():
    parser = argparse.ArgumentParser(description="MySQL 数据导出到 Excel 工具")
    parser.add_argument("-H", "--host", default="localhost", help="MySQL 主机地址")
    parser.add_argument("-P", "--port", type=int, default=3306, help="MySQL 端口")
    parser.add_argument("-u", "--user", default="root", help="MySQL 用户名")
    parser.add_argument("-p", "--password", default="", help="MySQL 密码")
    parser.add_argument("-d", "--database", required=True, help="数据库名")
    parser.add_argument("--docker", help="使用 Docker 容器（容器名称）")
    
    subparsers = parser.add_subparsers(dest='command', help='命令')
    
    # 导出表
    parser_table = subparsers.add_parser('table', help='导出表数据')
    parser_table.add_argument("table_name", help="表名")
    parser_table.add_argument("-o", "--output", help="输出 Excel 文件路径或目录")
    parser_table.add_argument("-w", "--where", default="", help="WHERE 条件（不包含 WHERE 关键字）")
    parser_table.add_argument("-l", "--limit", type=int, help="限制导出行数")
    parser_table.add_argument("-b", "--batch", action="store_true", help="分批导出（每批1000行）")
    parser_table.add_argument("--batch-size", type=int, default=1000, help="每批导出的行数（默认1000）")
    
    # 导出查询
    parser_query = subparsers.add_parser('query', help='导出查询结果')
    parser_query.add_argument("query", help="SQL 查询语句")
    parser_query.add_argument("-o", "--output", required=True, help="输出 Excel 文件路径")
    
    # 导出所有表
    parser_all = subparsers.add_parser('all', help='导出所有表')
    parser_all.add_argument("-o", "--output", help="输出 Excel 文件路径（所有表在一个文件中）")
    parser_all.add_argument("-d", "--output-dir", help="输出目录（每个表一个文件）")
    
    # 列出表
    subparsers.add_parser('list', help='列出所有表')
    
    # 导出 law_regulation 到 Word
    parser_word = subparsers.add_parser('law-word', help='导出 law_regulation 表到 Word 文件（每行一个文件）')
    parser_word.add_argument("-o", "--output", help="输出目录")
    parser_word.add_argument("-t", "--table", default="law_regulation", help="表名（默认: law_regulation）")
    
    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        return
    
    use_docker = args.docker is not None
    container_name = args.docker or "law-manage-mysql"
    
    exporter = MySQLToExcelExporter(
        host=args.host,
        port=args.port,
        user=args.user,
        password=args.password,
        database=args.database,
        use_docker=use_docker,
        container_name=container_name
    )
    
    if args.command == 'table':
        exporter.export_table(
            args.table_name,
            args.output,
            where_clause=args.where,
            limit=args.limit,
            batch_export=args.batch,
            batch_size=args.batch_size
        )
    elif args.command == 'query':
        exporter.export_query(args.query, args.output)
    elif args.command == 'all':
        exporter.export_all_tables(args.output, args.output_dir)
    elif args.command == 'list':
        exporter.list_tables()
    elif args.command == 'law-word':
        exporter.export_law_regulation_to_word(args.output, args.table)


if __name__ == "__main__":
    if len(sys.argv) == 1:
        # 交互式模式
        print("=== MySQL 数据导出到 Excel 工具 ===\n")
        
        use_docker = input("是否使用 Docker 容器？(y/N): ").strip().lower() == 'y'
        
        if use_docker:
            container_name = input("容器名称 (默认: law-manage-mysql): ").strip() or "law-manage-mysql"
            host = "localhost"
            port = 3306
        else:
            host = input("MySQL 主机 (默认: localhost): ").strip() or "localhost"
            port = int(input("MySQL 端口 (默认: 3306): ").strip() or "3306")
            container_name = None
        
        user = input("MySQL 用户名 (默认: root): ").strip() or "root"
        password = input("MySQL 密码: ").strip()
        database = input("数据库名: ").strip()
        
        exporter = MySQLToExcelExporter(
            host=host,
            port=port,
            user=user,
            password=password,
            database=database,
            use_docker=use_docker,
            container_name=container_name
        )
        
        print("\n请选择操作：")
        print("1. 列出所有表")
        print("2. 导出单个表")
        print("3. 导出查询结果")
        print("4. 导出所有表（一个文件）")
        print("5. 导出所有表（每个表一个文件）")
        print("6. 导出 law_regulation 到 Word（每行一个文件）")
        
        choice = input("\n请输入选项 (1-6): ").strip()
        
        if choice == "1":
            exporter.list_tables()
        elif choice == "2":
            table_name = input("表名: ").strip()
            batch_export = input("是否分批导出（每批1000行）？(y/N): ").strip().lower() == 'y'
            
            if batch_export:
                batch_size = input("每批行数（默认1000）: ").strip()
                batch_size = int(batch_size) if batch_size else 1000
                output_dir = input("输出目录（留空自动生成）: ").strip() or None
                where_clause = input("WHERE 条件（可选，留空则导出全部）: ").strip()
                limit = input("限制总行数（可选，留空则导出全部）: ").strip()
                limit = int(limit) if limit else None
                exporter.export_table(
                    table_name, output_dir, 
                    where_clause=where_clause, 
                    limit=limit,
                    batch_export=True,
                    batch_size=batch_size
                )
            else:
                output_file = input("输出文件路径（留空自动生成）: ").strip() or None
                where_clause = input("WHERE 条件（可选，留空则导出全部）: ").strip()
                limit = input("限制行数（可选，留空则不限制）: ").strip()
                limit = int(limit) if limit else None
                exporter.export_table(table_name, output_file, where_clause=where_clause, limit=limit)
        elif choice == "3":
            query = input("SQL 查询语句: ").strip()
            output_file = input("输出文件路径: ").strip()
            exporter.export_query(query, output_file)
        elif choice == "4":
            output_file = input("输出文件路径: ").strip()
            exporter.export_all_tables(output_file=output_file)
        elif choice == "5":
            output_dir = input("输出目录: ").strip()
            exporter.export_all_tables(output_dir=output_dir)
        elif choice == "6":
            table_name = input("表名 (默认: law_regulation): ").strip() or "law_regulation"
            output_dir = input("输出目录（留空自动生成）: ").strip() or None
            exporter.export_law_regulation_to_word(output_dir, table_name)
        else:
            print("无效选项")
    else:
        main()

